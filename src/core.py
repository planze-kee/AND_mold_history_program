"""
통합 핵심 모듈: HWP 처리, 문서 생성, 이미지 처리
"""
import argparse
import csv
import hashlib
import json
import logging
import re
import shutil
import zlib
from dataclasses import dataclass, field
import struct
import olefile
from datetime import datetime, date
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Set, Tuple
from io import BytesIO

from docx import Document
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from PIL import Image
import sys

logger = logging.getLogger(__name__)

from .constants import HWPConstants, DocumentConstants, PathConstants, HWPFormConstants


# ============================================================================
# 사용자 정의 예외
# ============================================================================
class HWPProcessingError(Exception):
    """HWP 파일 처리 중 발생하는 오류"""
    pass


class ImageExtractionError(Exception):
    """이미지 추출 중 발생하는 오류"""
    pass


class DocumentGenerationError(Exception):
    """Word 문서 생성 중 발생하는 오류"""
    pass


class XLSXProcessingError(Exception):
    """XLSX 파일 처리 중 발생하는 오류"""
    pass


# ============================================================================
# HWP 텍스트 추출
# ============================================================================
class HWPTextExtractor:
    """Extract text from HWP5 BodyText stream"""

    PARA_TEXT = HWPConstants.PARA_TEXT_TAG

    def __init__(self, section_data: bytes):
        self.section_data = section_data

    def extract_all_text(self) -> List[str]:
        """Extract all text from decompressed section"""
        texts = []
        pos = 0

        while pos < len(self.section_data):
            if pos + 4 > len(self.section_data):
                break

            # Parse record header (4 bytes)
            hdr = struct.unpack('<I', self.section_data[pos:pos+4])[0]
            tag = hdr & 0x3FF
            level = (hdr >> 10) & 0x3FF
            size = (hdr >> 20) & 0xFFF

            pos += 4

            # Handle extended size
            if size == 0xFFF:
                if pos + 4 > len(self.section_data):
                    break
                size = struct.unpack('<I', self.section_data[pos:pos+4])[0]
                pos += 4

            if pos + size > len(self.section_data):
                break

            payload = self.section_data[pos:pos+size]

            # Extract PARA_TEXT records (tag 67)
            if tag == self.PARA_TEXT and size > 0:
                text = self._decode_text(payload)
                if text.strip():
                    texts.append(text)

            pos += size

        return texts

    def _decode_text(self, payload: bytes) -> str:
        """Decode text from PARA_TEXT payload (UTF-16LE)"""
        result = []

        for i in range(0, len(payload), 2):
            if i + 1 < len(payload):
                char_code = struct.unpack('<H', payload[i:i+2])[0]

                if char_code == 0x000D:
                    break
                elif char_code < 0x0020 and char_code != 0:
                    continue
                else:
                    try:
                        result.append(chr(char_code))
                    except (ValueError, OverflowError):
                        pass

        return ''.join(result).strip()


# ============================================================================
# HWP 필드 추출기 (HWPDataExtractor 내부 로직 분리)
# ============================================================================
class HWPFieldExtractor:
    """HWP 폼 텍스트 리스트에서 필드별 값을 추출하는 클래스.

    _extract_fields()의 복잡한 파싱 로직을 역할별 메서드로 분리하여
    가독성과 테스트 가능성을 높인다.
    """

    def __init__(self, form_texts: List[str]):
        self.texts = form_texts
        self.boundary = HWPFormConstants.BOUNDARY_LABELS
        self.lookahead = HWPFormConstants.FIELD_LOOKAHEAD
        self._size_pattern = re.compile(HWPConstants.__dict__.get('SIZE_PATTERN', r'\d{1,5}x\d{1,5}x\d{1,5}'), re.IGNORECASE)

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------
    def extract(self, row: Dict[str, str]) -> None:
        """폼 텍스트 전체에서 모든 필드를 추출하여 row에 기록"""
        self._apply_pattern1(row)
        self._apply_pattern2(row)

    # ------------------------------------------------------------------
    # Pattern 1: "라벨 : 값" 동일 라인 형식
    # ------------------------------------------------------------------
    def _apply_pattern1(self, row: Dict[str, str]) -> None:
        for text in self.texts:
            if ':' not in text:
                continue
            parts = text.split(':', 1)
            if len(parts) != 2:
                continue
            label, value = parts[0].strip(), parts[1].strip()
            for key_label, row_key in HWPFormConstants.INLINE_FIELD_MAP.items():
                if key_label in label:
                    row[row_key] = value
                    break

    # ------------------------------------------------------------------
    # Pattern 2: 라벨 다음 줄에 값이 오는 형식
    # ------------------------------------------------------------------
    def _apply_pattern2(self, row: Dict[str, str]) -> None:
        i = 0
        while i < len(self.texts):
            text = self.texts[i]

            # 일반 다음줄 필드
            if text in HWPFormConstants.NEXTLINE_FIELD_MAP:
                row_key = HWPFormConstants.NEXTLINE_FIELD_MAP[text]
                val = self._find_simple_next(i)
                if val:
                    row[row_key] = val

            # 경계에서 break 처리하는 필드 (GATE/機械/契約日)
            elif text in HWPFormConstants.BREAK_ON_BOUNDARY_FIELD_MAP:
                row_key = HWPFormConstants.BREAK_ON_BOUNDARY_FIELD_MAP[text]
                val = self._find_next_with_break(i)
                if val:
                    row[row_key] = val

            # 承認日 — 年 포함 값만 수용
            elif text == HWPFormConstants.APPROVAL_DATE_LABEL:
                val = self._find_approval_date(i)
                if val:
                    row[HWPFormConstants.APPROVAL_DATE_KEY] = val

            # 金型規格 — XxXxX 패턴 검사
            elif text == HWPFormConstants.MOLD_SIZE_LABEL or text == HWPFormConstants.MOLD_SIZE_ALT_LABEL:
                row[HWPFormConstants.MOLD_SIZE_KEY] = self._extract_mold_size(i)

            # 金型材質 — BASE / CORE 추출
            elif text == HWPFormConstants.MOLD_MATERIAL_LABEL:
                base_val, core_val = self._extract_mold_material(i)
                row[HWPFormConstants.MOLD_MATERIAL_BASE_KEY] = base_val
                row[HWPFormConstants.MOLD_MATERIAL_CORE_KEY] = core_val

            i += 1

    # ------------------------------------------------------------------
    # 헬퍼: 다음 줄 값 탐색
    # ------------------------------------------------------------------
    def _find_simple_next(self, label_idx: int) -> str:
        """경계 라벨이 아닌 첫 번째 비어있지 않은 값을 반환"""
        for j in range(label_idx + 1, min(label_idx + self.lookahead, len(self.texts))):
            val = self.texts[j]
            if val.strip() and val not in self.boundary:
                return val
        return ""

    def _find_next_with_break(self, label_idx: int) -> str:
        """경계 라벨에서 즉시 탐색 중단 (값이 없으면 빈 값)"""
        for j in range(label_idx + 1, min(label_idx + self.lookahead, len(self.texts))):
            val = self.texts[j]
            if val in self.boundary:
                break
            if val.strip():
                return val
        return ""

    def _find_approval_date(self, label_idx: int) -> str:
        """承認日: '年' 문자가 포함된 값을 탐색"""
        for j in range(label_idx + 1, min(label_idx + self.lookahead, len(self.texts))):
            val = self.texts[j]
            if val.strip() and HWPFormConstants.APPROVAL_DATE_MARKER in val:
                return val
        return ""

    def _extract_mold_size(self, label_idx: int) -> str:
        """金型規格: XxXxX 패턴 탐색"""
        for j in range(label_idx + 1, min(label_idx + HWPFormConstants.MOLD_SIZE_LOOKAHEAD, len(self.texts))):
            if self._size_pattern.search(self.texts[j]):
                return self.texts[j]
        return ""

    def _extract_mold_material(self, label_idx: int) -> Tuple[str, str]:
        """金型材質: BASE / CORE 값을 가로 배치에서 추출"""
        base_value = ""
        core_value = ""
        j = label_idx + 1
        limit = min(label_idx + HWPFormConstants.MATERIAL_LOOKAHEAD, len(self.texts))

        while j < limit:
            current = self.texts[j]

            if current == 'BASE':
                for k in range(j + 1, len(self.texts)):
                    nxt = self.texts[k]
                    if nxt == 'CORE':
                        break
                    if nxt.strip() and nxt not in self.boundary:
                        base_value = nxt
                        break

            elif current == 'CORE':
                for k in range(j + 1, len(self.texts)):
                    nxt = self.texts[k]
                    if nxt in self.boundary and nxt not in ['BASE', 'CORE']:
                        break
                    if nxt.strip() and nxt not in self.boundary:
                        core_value = nxt
                        break

            j += 1

        return base_value, core_value


# ============================================================================
# HWP 데이터 추출
# ============================================================================
class HWPDataExtractor:
    """Extract form data from HWP using olefile"""

    def __init__(self, hwp_filepath: str):
        self.filepath = hwp_filepath
        self.filename = Path(hwp_filepath).stem

    def sanitize_filename(self, text: str) -> str:
        """파일명으로 사용할 수 없는 문자를 '_'로 치환"""
        return re.sub(DocumentConstants.INVALID_FILENAME_CHARS, DocumentConstants.INVALID_FILENAME_REPLACEMENT, text)

    def extract(self) -> Dict[str, str]:
        """Extract all form fields from HWP file"""
        row = {chr(65+i): '' for i in range(HWPConstants.FIELD_COUNT)}  # A-AC
        row['A'] = self.filename

        try:
            ole = olefile.OleFileIO(self.filepath)
            all_streams = ole.listdir()

            section_data = None
            for stream_path in all_streams:
                stream_str = '/'.join(stream_path) if isinstance(stream_path, list) else str(stream_path)
                if 'BodyText' in stream_str and 'Section0' in stream_str:
                    try:
                        section_data = ole.openstream(stream_path).read()
                        break
                    except Exception:
                        pass

            if section_data is None:
                ole.close()
                return row

            try:
                decompressed = zlib.decompress(section_data, -15)
            except zlib.error:
                decompressed = section_data

            extractor = HWPTextExtractor(decompressed)
            texts = extractor.extract_all_text()
            self._extract_fields(row, texts)

            ole.close()

        except Exception as e:
            logger.warning(f"HWP 데이터 추출 실패: {self.filepath} - {e}")

        return row

    def _extract_fields(self, row: Dict[str, str], texts: List[str]) -> None:
        """Extract field values from text list with form-specific parsing.

        폼 시작 위치를 찾아 HWPFieldExtractor에 위임한다.
        """
        form_start = 0
        for i, text in enumerate(texts):
            if HWPFormConstants.FORM_START_LABEL in text and ':' in text:
                form_start = i
                break

        if form_start >= len(texts):
            return

        form_texts = texts[form_start:]
        extractor = HWPFieldExtractor(form_texts)
        extractor.extract(row)

        # 금형사진 파일명 자동 생성 및 저장 (품명_도면번호 형식)
        product_name = row['J'].strip()
        drawing_no = row['K'].strip()

        if product_name or drawing_no:
            product_name_safe = self.sanitize_filename(product_name)
            drawing_no_safe = self.sanitize_filename(drawing_no)

            image_filename = f"{product_name_safe}_{drawing_no_safe}"
            row[chr(65 + HWPFormConstants.IMAGE_FIELD_INDEX)] = image_filename


# ============================================================================
# HWP 처리기 (HWP → XLSX)
# ============================================================================

def _extract_hwp_worker(hwp_path: Path) -> Tuple[Path, List[str]]:
    """멀티프로세싱 워커: HWP 파일 하나를 추출하고 (파일경로, 행값) 반환.

    ProcessPoolExecutor에서 실행되므로 최상위(top-level) 함수여야 한다.
    """
    extractor = HWPDataExtractor(str(hwp_path))
    row_data = extractor.extract()
    row_values = [row_data.get(chr(65 + i), "") for i in range(HWPConstants.FIELD_COUNT)]
    return hwp_path, row_values


class HWPProcessor:
    """Process HWP files and convert to XLSX"""

    CANONICAL_HEADERS = DocumentConstants.CANONICAL_HEADERS

    # 멀티프로세싱 활성화 기준 파일 수 (이 수 이상이면 병렬 처리)
    MP_THRESHOLD: int = 5

    @classmethod
    def extract_rows_from_hwp(
        cls,
        input_dir: Path,
        callback: Optional[Callable[[str], None]] = None,
        workers: Optional[int] = None,
    ) -> List[List[str]]:
        """HWP 파일들을 읽어 행 목록을 반환한다.

        Args:
            input_dir: HWP 파일이 있는 디렉터리.
            callback: 진행 상황을 전달할 함수 (GUI 연동용).
            workers: 병렬 프로세스 수. None이면 CPU 수에 따라 자동 결정.
                     파일 수가 MP_THRESHOLD 미만이면 단일 프로세스로 실행.
        """
        import multiprocessing
        from concurrent.futures import ProcessPoolExecutor, as_completed

        hwp_files = sorted(input_dir.rglob("*.hwp"))
        if not hwp_files:
            raise FileNotFoundError(f"No HWP files found in: {input_dir}")

        total = len(hwp_files)

        # EXE(frozen) 환경에서는 ProcessPoolExecutor가 crash를 일으키므로 비활성화
        is_frozen = getattr(sys, "frozen", False)
        use_mp = total >= cls.MP_THRESHOLD and not is_frozen
        if use_mp:
            cpu_count = multiprocessing.cpu_count()
            max_workers = workers if workers else max(1, min(cpu_count, total))
        else:
            max_workers = 1

        # 순서 보장을 위해 결과를 {Path: row_values} 딕셔너리에 수집
        results: Dict[Path, List[str]] = {}

        if use_mp and max_workers > 1:
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                future_to_path = {
                    executor.submit(_extract_hwp_worker, f): f for f in hwp_files
                }
                done = 0
                for future in as_completed(future_to_path):
                    done += 1
                    hwp_path, row_values = future.result()
                    results[hwp_path] = row_values
                    msg = f"[{done}/{total}] 완료: {hwp_path.name}"
                    if callback:
                        callback(msg)
                    else:
                        print(msg)
        else:
            for idx, hwp_file in enumerate(hwp_files, start=1):
                msg = f"[{idx}/{total}] 추출 중: {hwp_file.name}"
                if callback:
                    callback(msg)
                else:
                    print(msg)
                _, row_values = _extract_hwp_worker(hwp_file)
                results[hwp_file] = row_values

        # 원래 파일 순서대로 정렬하여 반환
        return [results[f] for f in hwp_files]

    @classmethod
    def save_xlsx(cls, rows: List[List[str]], output_xlsx: Path) -> None:
        wb = Workbook()
        ws = wb.active
        ws.title = "data"

        ws.append(cls.CANONICAL_HEADERS)
        for row in rows:
            padded = list(row) + [""] * max(0, len(cls.CANONICAL_HEADERS) - len(row))
            ws.append(padded[:len(cls.CANONICAL_HEADERS)])

        output_xlsx.parent.mkdir(parents=True, exist_ok=True)
        wb.save(output_xlsx)

    @classmethod
    def process(
        cls,
        input_dir: Path,
        output_xlsx: Path,
        callback: Optional[Callable[[str], None]] = None,
        workers: Optional[int] = None,
    ) -> None:
        """Main entry point: HWP folder → XLSX file"""
        if not input_dir.exists():
            raise FileNotFoundError(f"input-dir not found: {input_dir}")

        rows = cls.extract_rows_from_hwp(input_dir, callback=callback, workers=workers)
        cls.save_xlsx(rows, output_xlsx)

        msg = f"✓ 저장 완료: {output_xlsx} ({len(rows)}행)"
        if callback:
            callback(msg)
        else:
            print(msg)


# ============================================================================
# 이미지 추출기
# ============================================================================
class HWPImageExtractor:
    """Extract images from HWP files"""

    MAGIC_BYTES = HWPConstants.IMAGE_MAGIC_BYTES

    def __init__(self, hwp_path: str, output_dir: str = 'img') -> None:
        self.hwp_path = hwp_path
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.file_name = Path(hwp_path).stem

        # HWP에서 품명과 도면번호 추출 및 파일명으로 안전한 형식으로 변환
        self.product_name = ""
        self.drawing_no = ""
        try:
            extractor = HWPDataExtractor(hwp_path)
            row_data = extractor.extract()
            self.product_name = self.sanitize_filename(row_data.get('J', '').strip())
            self.drawing_no = self.sanitize_filename(row_data.get('K', '').strip())
        except Exception as e:
            logger.warning(f"HWP 메타데이터 추출 실패 (파일명으로 대체): {hwp_path} - {e}")

    def sanitize_filename(self, text: str) -> str:
        """파일명으로 사용할 수 없는 문자를 '_'로 치환"""
        return re.sub(DocumentConstants.INVALID_FILENAME_CHARS, DocumentConstants.INVALID_FILENAME_REPLACEMENT, text)

    def detect_image_format(self, data: bytes) -> Optional[str]:
        for magic, fmt in self.MAGIC_BYTES.items():
            if data.startswith(magic):
                return fmt
        return None

    def try_decompress(self, data: bytes) -> bytes:
        """Try to decompress zlib compressed data"""
        try:
            return zlib.decompress(data, -15)
        except zlib.error:
            try:
                return zlib.decompress(data)
            except zlib.error:
                return data

    def try_fix_image(self, data: bytes) -> Tuple[Optional[bytes], Optional[str]]:
        try:
            img = Image.open(BytesIO(data))
            img.verify()
            return data, None
        except Exception:
            try:
                img = Image.open(BytesIO(data))
                output = BytesIO()
                img.convert('RGB').save(output, format='JPEG', quality=95)
                return output.getvalue(), 'jpg'
            except Exception as e:
                logger.debug(f"이미지 변환 실패 (스킵): {e}")
                return None, 'error'

    def extract_images(self) -> int:
        if not Path(self.hwp_path).exists():
            return 0

        if not olefile.isOleFile(self.hwp_path):
            return 0

        try:
            ole = olefile.OleFileIO(self.hwp_path)
        except (olefile.Error, OSError) as e:
            logger.warning(f"OLE 파일 열기 실패: {self.hwp_path} - {e}")
            return 0

        extracted_count = 0

        try:
            for entry in ole.listdir():
                entry_name = '/'.join(entry)

                if 'BinData' in entry_name and ole.get_type(entry) == olefile.STGTY_STREAM:
                    try:
                        raw_data = ole.openstream(entry).read()

                        if not raw_data:
                            continue

                        image_data = self.try_decompress(raw_data)

                        if not image_data:
                            continue

                        detected_fmt = self.detect_image_format(image_data)
                        fixed_data, conversion_fmt = self.try_fix_image(image_data)

                        if fixed_data is None:
                            continue

                        final_fmt = conversion_fmt or detected_fmt or 'jpg'
                        if final_fmt == 'pdf':
                            final_fmt = 'jpg'

                        image_index = extracted_count + 1
                        # 품명_도번 형식으로 생성 (예: VFD HOLDER_1072001450.jpg)
                        # 도번이 없으면 품명만으로 저장 (trailing underscore 없이)
                        if self.product_name and self.drawing_no:
                            base_filename = f"{self.product_name}_{self.drawing_no}"
                        elif self.product_name:
                            base_filename = self.product_name
                        else:
                            base_filename = self.file_name
                        if image_index == 1:
                            output_filename = f"{base_filename}.{final_fmt}"
                        else:
                            output_filename = f"{base_filename}_{image_index}.{final_fmt}"

                        output_path = self.output_dir / output_filename

                        with open(output_path, 'wb') as f:
                            f.write(fixed_data)

                        extracted_count += 1

                    except Exception as e:
                        logger.debug(f"BinData 스트림 처리 실패 (스킵): {e}")
                        continue

            ole.close()

        except Exception as e:
            logger.error(f"이미지 추출 중 오류: {self.hwp_path} - {e}")
            return 0

        return extracted_count


# ============================================================================
# 이미지 캐시 (반복 glob 방지)
# ============================================================================
class ImageCache:
    """디렉터리 내 이미지 파일을 한 번 인덱싱하여 O(1) 탐색을 제공한다.

    DocumentFiller.process() / DocxSyncManager.sync() 처럼 같은 img_dir를
    여러 행에 걸쳐 반복 탐색하는 경우 이 캐시를 먼저 생성하고 넘겨준다.
    이미지 파일이 추가/삭제된 경우 invalidate() 후 rebuild()를 호출한다.
    """

    def __init__(self, img_dir: Path) -> None:
        self.img_dir = img_dir
        self._index: Dict[str, Path] = {}
        if img_dir.exists():
            self._build()

    # ------------------------------------------------------------------
    # 인덱스 빌드
    # ------------------------------------------------------------------
    def _build(self) -> None:
        """이미지 디렉터리 전체를 스캔하여 stem → Path 인덱스를 구성한다.

        - 정확한 stem: "VFD HOLDER_1072001450"
        - sanitized stem: "VFD HOLDER_1072001450" (특수문자 치환)
        - 연번 제거 stem: "VFD HOLDER_1072001450" (말미 _2, _3 ... 제거)
        """
        self._index.clear()
        exts = set(DocumentConstants.IMAGE_EXTS)
        for p in self.img_dir.iterdir():
            if not p.is_file() or p.suffix.lower() not in exts:
                continue
            stem = p.stem

            # 정확한 stem 등록 (나중에 들어온 파일이 앞의 것을 덮어쓰지 않도록
            # 연번 없는 파일을 우선)
            self._register(stem, p)

            # sanitized stem (특수문자 → '_')
            sanitized = re.sub(
                DocumentConstants.INVALID_FILENAME_CHARS,
                DocumentConstants.INVALID_FILENAME_REPLACEMENT,
                stem,
            )
            if sanitized != stem:
                self._register(sanitized, p)

            # 연번 접미사 제거 stem (예: "xxx_2" → "xxx")
            base = re.sub(r'_\d+$', '', stem)
            if base != stem:
                self._register(base, p)

    def _register(self, key: str, path: Path) -> None:
        """연번 없는 파일(우선순위 높음)이 있을 때 덮어쓰지 않는다."""
        if key not in self._index:
            self._index[key] = path
        else:
            # 현재 등록된 파일이 연번 접미사를 가지면 새 파일로 교체
            existing_stem = self._index[key].stem
            if re.search(r'_\d+$', existing_stem):
                self._index[key] = path

    # ------------------------------------------------------------------
    # 탐색
    # ------------------------------------------------------------------
    def find(self, stem: str) -> Optional[Path]:
        """stem(확장자 없음)으로 이미지를 O(1)에 검색한다.

        1. 정확한 stem 탐색
        2. sanitized stem 탐색
        3. img_dir가 바뀐 경우를 위해 존재 여부 재확인
        """
        # 1. 정확한 stem
        result = self._index.get(stem)
        if result and result.exists():
            return result

        # 2. sanitized stem
        sanitized = re.sub(
            DocumentConstants.INVALID_FILENAME_CHARS,
            DocumentConstants.INVALID_FILENAME_REPLACEMENT,
            stem,
        )
        if sanitized != stem:
            result = self._index.get(sanitized)
            if result and result.exists():
                return result

        # 3. trailing underscore/dash 제거 후 재시도
        clean = stem.rstrip('_- ')
        if clean and clean != stem:
            result = self._index.get(clean)
            if result and result.exists():
                return result

        return None

    def invalidate(self) -> None:
        """캐시 무효화 후 재빌드 (파일 추가/삭제 후 호출)"""
        if self.img_dir.exists():
            self._build()
        else:
            self._index.clear()

    def __len__(self) -> int:
        return len(self._index)


# ============================================================================
# 문서 채우기 (XLSX → DOCX)
# ============================================================================
class DocumentFiller:
    """Fill DOCX from XLSX with placeholder tokens and images"""

    CHECKBOX_LABELS = DocumentConstants.CHECKBOX_LABELS
    IMAGE_TOKEN = DocumentConstants.IMAGE_TOKEN
    IMAGE_EXTS = DocumentConstants.IMAGE_EXTS
    ALIASES = DocumentConstants.ALIASES

    @staticmethod
    def normalize(text: str) -> str:
        if text is None:
            return ""
        return re.sub(r"\s+", "", str(text)).strip().lower()

    @classmethod
    def row_norm_map(cls, row: Dict[str, str]) -> Dict[str, str]:
        return {cls.normalize(k): (v or "").strip() for k, v in row.items()}

    @classmethod
    def value_by_aliases(cls, norm_row: Dict[str, str], aliases: List[str]) -> str:
        for alias in aliases:
            key = cls.normalize(alias)
            if key in norm_row and norm_row[key] != "":
                return norm_row[key]
        return ""

    @classmethod
    def value_for_label(cls, label: str, norm_row: Dict[str, str]) -> str:
        if label == "金型写真":
            return cls.IMAGE_TOKEN

        if label in cls.CHECKBOX_LABELS:
            raw = cls.value_by_aliases(norm_row, [label] + cls.ALIASES.get(label, []))
            return "O" if raw == "1" else ""

        direct = cls.value_by_aliases(norm_row, [label])
        if direct != "":
            return direct

        if label in cls.ALIASES:
            return cls.value_by_aliases(norm_row, cls.ALIASES[label])

        return ""

    @classmethod
    def replace_placeholders(cls, text: str, norm_row: Dict[str, str]) -> Tuple[str, int, Set[str]]:
        if not text:
            return text, 0, set()

        replaced_count = 0
        replaced_labels: Set[str] = set()

        def repl(match):
            nonlocal replaced_count
            label = match.group(0)[1:-1].strip()
            value = cls.value_for_label(label, norm_row)
            replaced_count += 1
            replaced_labels.add(label)
            return value

        return re.sub(r"\{[^{}]+\}", repl, text), replaced_count, replaced_labels

    @staticmethod
    def apply_small_font_if_needed(paragraph: Any, replaced_labels: Set[str]) -> None:
        if not ({"BASE", "CORE"} & replaced_labels):
            return
        for run in paragraph.runs:
            run.font.size = Pt(DocumentConstants.MOLD_MATERIAL_FONT_SIZE_PT)

    @classmethod
    def _copy_text_run_format(cls, src_run: Any, dst_run: Any) -> None:
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        dst_run.style = src_run.style
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size

    @classmethod
    def _insert_run_after(cls, paragraph: Any, after_run: Any, text: str = "") -> Any:
        new_run = paragraph.add_run(text)
        after_run._r.addnext(new_run._r)
        return new_run

    @classmethod
    def _insert_image_in_paragraph(cls, paragraph: Any, image_path: Optional[Path]) -> int:
        para_text = paragraph.text or ""
        if cls.IMAGE_TOKEN not in para_text:
            return 0

        inserted = 0
        original_runs = list(paragraph.runs)

        # Case 1: token exists fully inside at least one run
        token_in_single_run = False
        for run in original_runs:
            run_text = run.text or ""
            if cls.IMAGE_TOKEN not in run_text:
                continue
            token_in_single_run = True

            parts = run_text.split(cls.IMAGE_TOKEN)
            run.text = parts[0]
            cursor = run

            for part in parts[1:]:
                if image_path and image_path.exists():
                    img_run = cls._insert_run_after(paragraph, cursor)
                    img_run.add_picture(str(image_path), width=Cm(DocumentConstants.IMAGE_WIDTH_CM), height=Cm(DocumentConstants.IMAGE_HEIGHT_CM))
                    cursor = img_run
                    inserted += 1

                if part:
                    txt_run = cls._insert_run_after(paragraph, cursor, part)
                    cls._copy_text_run_format(run, txt_run)
                    cursor = txt_run

        if token_in_single_run:
            return inserted

        # Case 2: token is split across multiple runs
        if not original_runs:
            return 0

        base_run = original_runs[0]
        for r in original_runs:
            r.text = ""

        parts = para_text.split(cls.IMAGE_TOKEN)
        base_run.text = parts[0]
        cursor = base_run

        for part in parts[1:]:
            if image_path and image_path.exists():
                img_run = cls._insert_run_after(paragraph, cursor)
                img_run.add_picture(str(image_path), width=Cm(DocumentConstants.IMAGE_WIDTH_CM), height=Cm(DocumentConstants.IMAGE_HEIGHT_CM))
                cursor = img_run
                inserted += 1

            if part:
                txt_run = cls._insert_run_after(paragraph, cursor, part)
                cls._copy_text_run_format(base_run, txt_run)
                cursor = txt_run

        return inserted

    @classmethod
    def insert_images_by_token(cls, doc: Document, image_path: Optional[Path]) -> int:
        count = 0
        for p in doc.paragraphs:
            count += cls._insert_image_in_paragraph(p, image_path)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        count += cls._insert_image_in_paragraph(p, image_path)
        return count

    @classmethod
    def find_image_for_output(cls, img_dir: Path, output_stem: str) -> Optional[Path]:
        """이미지 파일 찾기 (품명_도번 형식)

        1. 정확한 파일명으로 먼저 찾기
        2. sanitize 처리된 파일명으로 찾기
        3. output_stem에 포함된 도번으로도 확인
        """
        if not output_stem or not img_dir.exists():
            return None

        # 정확한 파일명으로 찾기
        for ext in cls.IMAGE_EXTS:
            p = img_dir / f"{output_stem}{ext}"
            if p.exists():
                return p

        # sanitize 처리한 이름으로 찾기 (파일명 사용 불가 문자 처리)
        sanitized_stem = re.sub(DocumentConstants.INVALID_FILENAME_CHARS, DocumentConstants.INVALID_FILENAME_REPLACEMENT, output_stem)
        if sanitized_stem != output_stem:
            for ext in cls.IMAGE_EXTS:
                p = img_dir / f"{sanitized_stem}{ext}"
                if p.exists():
                    return p

        # output_stem을 sanitize한 후 glob으로 찾기
        # 예: output_stem="1072001446 / 2001447"
        #     sanitized="1072001446 _ 2001447"
        #     파일: "FRONT POLE _ REAR POLE_1072001446 _ 2001447.jpg"
        # → glob: "*1072001446 _ 2001447*"로 검색
        matches = []
        search_stem = sanitized_stem  # 이미 sanitize된 버전 사용
        for ext in cls.IMAGE_EXTS:
            try:
                # 파일명에 공백이 있을 수 있으니 escape 처리
                matches.extend(img_dir.glob(f"*{search_stem}*{ext}"))
            except (OSError, ValueError):
                pass

        # trailing underscore/hyphen 제거 후 재시도
        # 예: XLSX에 "MR14 CASE_"로 저장된 경우 → "MR14 CASE"로 검색
        clean_stem = search_stem.rstrip('_- ')
        if clean_stem and clean_stem != search_stem:
            for ext in cls.IMAGE_EXTS:
                try:
                    p = img_dir / f"{clean_stem}{ext}"
                    if p.exists():
                        return p
                    matches.extend(img_dir.glob(f"*{clean_stem}*{ext}"))
                except (OSError, ValueError):
                    pass

        if matches:
            # 연번이 없는 파일 우선 (예: "xxx_yyy.jpg" > "xxx_yyy_2.jpg")
            matches_no_suffix = [m for m in matches if not m.stem.endswith(DocumentConstants.IMAGE_NUMBERED_SUFFIXES)]
            if matches_no_suffix:
                return matches_no_suffix[0]
            return matches[0]

        return None

    @classmethod
    def replace_in_doc(cls, doc: Document, norm_row: Dict[str, str]) -> int:
        total = 0
        for p in doc.paragraphs:
            new_text, cnt, labels = cls.replace_placeholders(p.text, norm_row)
            if new_text != p.text:
                p.text = new_text
                cls.apply_small_font_if_needed(p, labels)
            total += cnt
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        new_text, cnt, labels = cls.replace_placeholders(p.text, norm_row)
                        if new_text != p.text:
                            p.text = new_text
                            cls.apply_small_font_if_needed(p, labels)
                        total += cnt
        return total

    @classmethod
    def pick_output_name(cls, row: Dict[str, str], idx: int) -> str:
        nrow = cls.row_norm_map(row)
        value = cls.value_by_aliases(nrow, ["File name", "file_name", "filename", "source_hwp"])
        if value:
            m = re.search(r"\d{2}-\d{3}", value)
            if m:
                return m.group(0)
            digits = re.findall(r"\d+", value)
            if digits:
                return "".join(digits)
            stem = Path(value).stem.strip()
            if stem:
                return stem
        return f"row_{idx:03d}"

    @staticmethod
    def unique_path(path: Path) -> Path:
        if not path.exists():
            return path
        n = 2
        while True:
            c = path.with_name(f"{path.stem}_{n}{path.suffix}")
            if not c.exists():
                return c
            n += 1

    @classmethod
    def load_rows_from_xlsx(cls, xlsx_path: Path) -> List[Dict[str, str]]:
        wb = load_workbook(xlsx_path, data_only=True)
        ws = wb.active
        headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]

        rows: List[Dict[str, str]] = []
        for row_values in ws.iter_rows(min_row=2, values_only=True):
            row_dict: Dict[str, str] = {}
            has_any = False
            for i, header in enumerate(headers):
                v = row_values[i] if i < len(row_values) else None
                text = "" if v is None else str(v).strip()
                if text:
                    has_any = True
                row_dict[header] = text
            if has_any:
                rows.append(row_dict)
        wb.close()
        return rows

    @classmethod
    def process(cls, xlsx_path: Path, template_path: Path, out_dir: Path, img_dir: Path,
                limit: int = 0, callback: Optional[Callable[[str], None]] = None) -> None:
        """Main entry point: XLSX + Template → DOCX files with images"""
        def _log(msg):
            if callback:
                callback(msg)
            else:
                print(msg)

        if not xlsx_path.exists():
            raise FileNotFoundError(f"XLSX not found: {xlsx_path}")
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        rows = cls.load_rows_from_xlsx(xlsx_path)
        if limit > 0:
            rows = rows[:limit]
        if not rows:
            _log("XLSX에 처리할 행이 없습니다.")
            return

        out_dir.mkdir(parents=True, exist_ok=True)

        # 이미지 파일명 저장 (XLSX 업데이트용)
        image_updates = {}  # {row_idx: image_filename}

        # ImageCache: img_dir 전체를 한 번 인덱싱 (행마다 반복 glob 방지)
        img_cache = ImageCache(img_dir)
        data_cache = ImageCache(out_dir / ".data")

        total = len(rows)
        for idx, row in enumerate(rows, start=1):
            doc = Document(str(template_path))
            nrow = cls.row_norm_map(row)
            replaced = cls.replace_in_doc(doc, nrow)

            out_name = cls.pick_output_name(row, idx)

            # 이미지 파일 찾기: 우선순위
            image_path = None

            # 1. XLSX의 "金型写真"(품명_도번)으로 먼저 시도
            image_name = row.get("金型写真", "").strip()
            if image_name:
                image_path = img_cache.find(image_name)

            # 2. "金型写真"으로 못 찾으면 XLSX의 도번으로 시도
            if not image_path:
                drawing_no = cls.value_by_aliases(nrow, ["図番番号", "drawing_no", "도번번호", "도번", "품번"])
                if drawing_no:
                    image_path = img_cache.find(drawing_no)

            # 2.5. 품명만으로 검색 (도번 없는 경우: XLSX에 "品名_"처럼 저장된 상황 대비)
            if not image_path:
                product_name = cls.value_by_aliases(nrow, ["品  名", "品 名", "product_name", "품명"])
                if product_name:
                    image_path = img_cache.find(product_name.strip())

            # 3. 도번으로도 못 찾으면 out_name(연번)으로 시도
            if not image_path:
                image_path = img_cache.find(out_name)

            # 4. img_dir에 없으면 out_dir/.data/ 도 검색 (신규 발행 첨부 이미지)
            if not image_path:
                search_name = image_name or out_name
                image_path = data_cache.find(search_name)

            inserted = cls.insert_images_by_token(doc, image_path)

            out_path = cls.unique_path(out_dir / f"{out_name}.docx")
            doc.save(str(out_path))

            # 실제 삽입된 이미지 파일명 저장 (XLSX 업데이트용)
            image_updates[idx] = image_path.name if image_path else ""

            img_status = "OK" if image_path else "NONE"
            _log(f"[{idx}/{total}] 저장: {out_name}.docx (치환={replaced}, 이미지={img_status})")

        # XLSX의 "金型写真" 컬럼 업데이트
        cls._update_xlsx_images(xlsx_path, image_updates)
        _log(f"✓ XLSX 이미지 컬럼 업데이트 완료")

    @classmethod
    def _update_xlsx_images(cls, xlsx_path: Path, image_updates: Dict[int, str]) -> None:
        """XLSX의 金型写真 컬럼에 실제 이미지 파일명 기록"""
        try:
            wb = load_workbook(xlsx_path)
            ws = wb.active

            # 金型写真 컬럼 정확하게 찾기 (전체 텍스트 매칭)
            photo_col_idx = None
            for col_idx in range(1, DocumentConstants.XLSX_MAX_COL_SEARCH):
                h = ws.cell(1, col_idx).value
                if h:
                    h_str = str(h).strip()
                    # 정확하게 "金型写真" 찾기 (다른 컬럼과 혼동 방지)
                    if "金型写真" in h_str:
                        photo_col_idx = col_idx
                        break

            # 金型写真 컬럼이 없으면 마지막 컬럼 다음에 추가
            if photo_col_idx is None:
                # 마지막 비어있지 않은 컬럼 찾기
                last_col = 1
                for col_idx in range(1, DocumentConstants.XLSX_MAX_COL_SEARCH):
                    if ws.cell(1, col_idx).value:
                        last_col = col_idx
                photo_col_idx = last_col + 1
                # 헤더 추가
                ws.cell(1, photo_col_idx).value = "金型写真"

            # 이미지 파일명 업데이트
            for row_idx, img_filename in image_updates.items():
                ws.cell(row_idx + 1, photo_col_idx).value = img_filename

            wb.save(xlsx_path)
            wb.close()
        except Exception as e:
            print(f"Warning: Could not update XLSX with image filenames: {e}")


# ============================================================================
# DocxSyncManager - XLSX-DOCX 동기화 (기능 1+2)
# ============================================================================
class DocxSyncManager:
    """XLSX 변경사항을 기존 DOCX 파일에 동기화 (파일명 변경 + 내용 갱신 + 이미지 rename)"""

    MANIFEST_FILENAME = "manifest.json"

    @classmethod
    def _manifest_path(cls, output_dir: Path) -> Path:
        return output_dir / cls.MANIFEST_FILENAME

    @classmethod
    def load_manifest(cls, output_dir: Path) -> Dict[str, Any]:
        """manifest.json 로드 (없으면 빈 딕셔너리)"""
        path = cls._manifest_path(output_dir)
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    @classmethod
    def save_manifest(cls, output_dir: Path, manifest: Dict[str, Any]) -> None:
        """manifest.json 저장"""
        output_dir.mkdir(parents=True, exist_ok=True)
        path = cls._manifest_path(output_dir)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)

    @classmethod
    def extract_serial(cls, file_name: str) -> str:
        """파일명에서 연번 추출 (예: '19-001' → '001')"""
        m = re.search(r"-(\d+)$", file_name)
        if m:
            return m.group(1)
        m = re.search(r"(\d+)$", file_name)
        if m:
            return m.group(1)
        return ""

    @classmethod
    def compute_row_hash(cls, row: Dict[str, Any]) -> str:
        """행 내용을 정규화하여 SHA1 해시 산출 (키 정렬로 순서 독립성 보장)"""
        norm = {k: str(v).strip() for k, v in sorted(row.items())}
        payload = json.dumps(norm, ensure_ascii=False, sort_keys=True)
        return hashlib.sha1(payload.encode("utf-8")).hexdigest()

    @classmethod
    def compute_image_sig(cls, image_path: Optional[Path]) -> str:
        """이미지 파일 서명 (경로 + mtime + size). 이미지 교체 감지용."""
        if not image_path or not image_path.exists():
            return ""
        st = image_path.stat()
        return f"{image_path.name}|{int(st.st_mtime)}|{st.st_size}"

    @classmethod
    def rename_image_files(cls, img_dir: Path, old_serial: str, new_serial: str, callback: Optional[Callable[[str], None]] = None) -> int:
        """연번 변경에 따른 이미지 파일명 변경"""
        if not img_dir.exists() or not old_serial or old_serial == new_serial:
            return 0
        renamed = 0
        for img_file in sorted(img_dir.iterdir()):
            if img_file.is_file() and img_file.name.startswith(f"{old_serial}_"):
                new_name = new_serial + img_file.name[len(old_serial):]
                new_path = img_dir / new_name
                if not new_path.exists():
                    img_file.rename(new_path)
                    renamed += 1
                    if callback:
                        callback(f"  이미지 rename: {img_file.name} → {new_name}")
        return renamed

    @classmethod
    def _resolve_image(
        cls,
        row: Dict[str, Any],
        out_name: str,
        img_cache: "ImageCache",
        data_cache: "ImageCache",
    ) -> Optional[Path]:
        """행 데이터에서 이미지 경로를 탐색하는 헬퍼 (sync/단건 공통)"""
        nrow = DocumentFiller.row_norm_map(row)
        image_name = row.get("金型写真", "").strip()
        image_path = None

        # 1. XLSX의 金型写真(품명_도번)으로 시도
        if image_name:
            image_path = img_cache.find(image_name)

        # 2. 図番番号로 시도
        if not image_path:
            drawing_no = DocumentFiller.value_by_aliases(
                nrow, ["図番番号", "drawing_no", "도번번호", "도번", "품번"])
            if drawing_no:
                image_path = img_cache.find(drawing_no)

        # 2.5. 品名으로 시도 (도번 없는 경우)
        if not image_path:
            product_name = DocumentFiller.value_by_aliases(
                nrow, ["品  名", "品 名", "product_name", "품명"])
            if product_name:
                image_path = img_cache.find(product_name.strip())

        # 3. out_name(연번)으로 시도
        if not image_path:
            image_path = img_cache.find(out_name)

        # 4. output_dir/.data/ 검색 (신규 발행 첨부 이미지)
        if not image_path:
            search_name = image_name or out_name
            image_path = data_cache.find(search_name)

        return image_path

    @classmethod
    def sync(cls, xlsx_path: Path, template_path: Path, output_dir: Path,
             img_dir: Path, callback: Optional[Callable[[str], None]] = None,
             force_all: bool = False) -> None:
        """XLSX 변경사항을 DOCX 파일에 동기화 (증분 재생성 지원)"""
        if callback:
            callback("동기화 시작...")

        rows = DocumentFiller.load_rows_from_xlsx(xlsx_path)
        manifest = cls.load_manifest(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # 템플릿 변경 자동 감지 → 변경 시 전체 강제 재생성
        current_tpl_hash = hashlib.sha1(template_path.read_bytes()).hexdigest()
        meta = manifest.get("_meta", {})
        if meta.get("template_hash") != current_tpl_hash:
            if not force_all and meta.get("template_hash"):
                # 이전 해시가 있는데 달라진 경우에만 알림 (최초 실행은 조용히)
                if callback:
                    callback("템플릿 변경 감지 → 전체 재생성")
            force_all = True

        # 1단계: 파일명 변경 감지 및 rename
        renamed_files = 0
        for idx, row in enumerate(rows, start=1):
            mgmt_no = row.get("管理番号", "").strip()
            if not mgmt_no or mgmt_no not in manifest:
                continue
            current_name = DocumentFiller.pick_output_name(row, idx)
            old_name = manifest[mgmt_no].get("file_name", "")
            if not old_name or old_name == current_name:
                continue

            # Word 파일 rename
            old_docx = output_dir / f"{old_name}.docx"
            new_docx = output_dir / f"{current_name}.docx"
            if old_docx.exists() and not new_docx.exists():
                old_docx.rename(new_docx)
                renamed_files += 1
                if callback:
                    callback(f"Word rename: {old_name}.docx → {current_name}.docx")

            # 이미지 파일 rename
            old_serial = cls.extract_serial(old_name)
            new_serial = cls.extract_serial(current_name)
            if old_serial and new_serial and old_serial != new_serial:
                img_count = cls.rename_image_files(img_dir, old_serial, new_serial, callback)
                if img_count > 0 and callback:
                    callback(f"  이미지 {img_count}개 rename 완료")

            # 이력 파일 rename (.data/ 하위 폴더)
            data_dir = output_dir / MaintenanceHistoryManager.DATA_DIR
            old_hist = data_dir / f"{old_name}_history.txt"
            new_hist = data_dir / f"{current_name}_history.txt"
            if old_hist.exists() and not new_hist.exists():
                old_hist.rename(new_hist)

        if renamed_files > 0 and callback:
            callback(f"파일명 변경: {renamed_files}건 처리됨")

        # 2단계: 증분 DOCX 재생성
        if callback:
            callback("DOCX 내용 갱신 중...")

        total = len(rows)
        new_manifest: Dict[str, Any] = {"_meta": {"template_hash": current_tpl_hash}}

        # ImageCache: 1단계 rename 이후 빌드해야 최신 파일명을 반영
        img_cache = ImageCache(img_dir)
        data_cache = ImageCache(output_dir / ".data")

        skipped = 0
        regenerated = 0
        created = 0

        for idx, row in enumerate(rows, start=1):
            mgmt_no = row.get("管理番号", "").strip()
            out_name = DocumentFiller.pick_output_name(row, idx)
            out_path = output_dir / f"{out_name}.docx"

            # 현재 상태 계산
            current_hash = cls.compute_row_hash(row)
            image_path = cls._resolve_image(row, out_name, img_cache, data_cache)
            current_img_sig = cls.compute_image_sig(image_path)

            # manifest 조회
            prev = manifest.get(mgmt_no, {})
            unchanged = (
                not force_all
                and out_path.exists()
                and prev.get("content_hash") == current_hash
                and prev.get("image_sig") == current_img_sig
            )

            if unchanged:
                skipped += 1
                new_manifest[mgmt_no] = prev
                if callback and idx % 50 == 0:
                    callback(f"[{idx}/{total}] 스킵 {skipped}건 / 재생성 {regenerated}건")
                continue

            try:
                doc = Document(str(template_path))
                nrow = DocumentFiller.row_norm_map(row)
                DocumentFiller.replace_in_doc(doc, nrow)
                DocumentFiller.insert_images_by_token(doc, image_path)
                doc.save(str(out_path))

                if mgmt_no in manifest and mgmt_no != "_meta":
                    regenerated += 1
                    tag = "갱신"
                else:
                    created += 1
                    tag = "신규"

                if mgmt_no:
                    new_manifest[mgmt_no] = {
                        "file_name": out_name,
                        "serial": cls.extract_serial(out_name),
                        "品名": row.get("品 名", "").strip(),
                        "図番番号": row.get("図番番号", "").strip(),
                        "content_hash": current_hash,
                        "image_sig": current_img_sig,
                        "last_updated": datetime.now().isoformat(timespec="seconds"),
                    }

                if callback:
                    callback(f"[{idx}/{total}] {tag}: {out_name}.docx")

            except Exception as e:
                if callback:
                    callback(f"[{idx}/{total}] 오류: {e}")

        cls.save_manifest(output_dir, new_manifest)
        if callback:
            callback(f"✓ 동기화 완료 — 신규 {created} / 갱신 {regenerated} / 스킵 {skipped} (manifest.json 저장됨)")


# ============================================================================
# MaintenanceHistoryManager - 유지보수 이력 관리 (기능 3)
# ============================================================================
class MaintenanceHistoryManager:
    """금형 유지보수 이력 관리 (.data/ 하위 폴더에 _history.txt 저장)"""

    DATA_DIR = ".data"

    @classmethod
    def get_history_path(cls, docx_path: Path) -> Path:
        data_dir = docx_path.parent / cls.DATA_DIR
        return data_dir / f"{docx_path.stem}_history.txt"

    @classmethod
    def read_history(cls, docx_path: Path) -> str:
        path = cls.get_history_path(docx_path)
        if path.exists():
            return path.read_text(encoding="utf-8")
        return f"# {docx_path.stem} 유지보수 이력\n\n"

    @classmethod
    def write_history(cls, docx_path: Path, content: str) -> None:
        path = cls.get_history_path(docx_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")

    @classmethod
    def add_entry(cls, docx_path: Path, entry_date: str, entry_type: str,
                  content: str, person: str) -> None:
        """새 유지보수 이력 항목 추가"""
        current = cls.read_history(docx_path)
        entry = (
            f"\n## {entry_date}\n"
            f"- 유형: {entry_type}\n"
            f"- 내용: {content}\n"
            f"- 담당자: {person}\n"
        )
        cls.write_history(docx_path, current.rstrip() + "\n" + entry)

    @classmethod
    def _summarize_history(cls, history_content: str) -> Tuple[str, str]:
        """이력 내용에서 修理内訳/事由 요약 텍스트 생성 (최근 3개 항목)"""
        lines = history_content.splitlines()
        entries: List[Dict] = []
        current: Optional[Dict] = None

        for line in lines:
            if line.startswith("## "):
                if current:
                    entries.append(current)
                current = {"date": line[3:].strip(), "type": "", "content": "", "person": ""}
            elif current:
                if line.startswith("- 유형:"):
                    current["type"] = line[5:].strip()
                elif line.startswith("- 내용:"):
                    current["content"] = line[5:].strip()
                elif line.startswith("- 담당자:"):
                    current["person"] = line[6:].strip()
        if current:
            entries.append(current)

        recent = entries[-3:]
        repair_parts = [f"{e['date']}: {e['content']}" for e in recent if e.get("content")]
        reason_parts = [e["type"] for e in recent if e.get("type")]

        return " / ".join(repair_parts), " / ".join(reason_parts)

    @classmethod
    def update_xlsx_reason(cls, docx_path: Path, xlsx_path: Path, content: str,
                           callback: Optional[Callable[[str], None]] = None) -> bool:
        """XLSX의 '事 由' 컬럼을 content로 갱신 (Word 재생성 없음)"""
        stem = docx_path.stem

        if not xlsx_path.exists():
            if callback:
                callback("오류: XLSX 파일을 찾을 수 없음")
            return False

        try:
            wb = load_workbook(xlsx_path)
            ws = wb.active
            headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]

            reason_col_idx = None
            file_name_col_idx = None
            for i, h in enumerate(headers):
                if h == "File name":
                    file_name_col_idx = i
                elif h == "事 由":
                    reason_col_idx = i

            if file_name_col_idx is None:
                wb.close()
                if callback:
                    callback("오류: XLSX에 'File name' 컬럼 없음")
                return False

            target_row_idx = None
            for row_idx, row_values in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                cell_val = str(row_values[file_name_col_idx] or "").strip()
                m = re.search(r"\d{2}-\d{3}", cell_val)
                row_stem = m.group(0) if m else Path(cell_val).stem
                if row_stem == stem:
                    target_row_idx = row_idx
                    break

            if target_row_idx is not None and reason_col_idx is not None:
                ws.cell(row=target_row_idx, column=reason_col_idx + 1).value = content

            wb.save(xlsx_path)
            wb.close()
            if callback:
                callback(f"XLSX 업데이트 완료 (行 {target_row_idx})")
            return True
        except PermissionError:
            if callback:
                callback(
                    f"XLSX 저장 실패: 파일이 다른 프로그램(Excel 등)에서 열려있습니다.\n"
                    f"  → {xlsx_path}\n"
                    f"  파일을 닫고 다시 시도하세요."
                )
            return False
        except Exception as e:
            if callback:
                callback(f"XLSX 업데이트 오류: {e}")
            return False

    @classmethod
    def apply_to_word(cls, docx_path: Path, xlsx_path: Path, template_path: Path,
                      img_dir: Path, callback: Optional[Callable[[str], None]] = None) -> bool:
        """이력 내용을 XLSX에 반영하고 Word 파일 재생성"""
        if not docx_path.exists():
            if callback:
                callback("오류: Word 파일을 찾을 수 없음")
            return False

        history_content = cls.read_history(docx_path)
        stem = docx_path.stem

        if not xlsx_path.exists():
            if callback:
                callback("오류: XLSX 파일을 찾을 수 없음")
            return False

        # XLSX에서 해당 행 찾아 事 由 업데이트
        ok = cls.update_xlsx_reason(docx_path, xlsx_path, history_content, callback)
        if not ok:
            return False

        # Word 파일 재생성 (해당 행 데이터로)
        if not template_path.exists():
            if callback:
                callback("오류: 템플릿 파일을 찾을 수 없음")
            return False

        try:
            rows = DocumentFiller.load_rows_from_xlsx(xlsx_path)
            target_row = None
            for idx, row in enumerate(rows, start=1):
                row_name = DocumentFiller.pick_output_name(row, idx)
                if row_name == stem:
                    target_row = row
                    break

            if target_row is None:
                if callback:
                    callback("오류: XLSX에서 해당 행을 찾을 수 없음")
                return False

            doc = Document(str(template_path))
            nrow = DocumentFiller.row_norm_map(target_row)
            DocumentFiller.replace_in_doc(doc, nrow)

            image_name = target_row.get("金型写真", "").strip()
            image_path = None

            # 1. XLSX의 金型写真(품명_도번)으로 시도
            if image_name:
                image_path = DocumentFiller.find_image_for_output(img_dir, image_name)

            # 2. 図番番号로 시도
            if not image_path:
                drawing_no = DocumentFiller.value_by_aliases(
                    nrow, ["図番番号", "drawing_no", "도번번호", "도번", "품번"])
                if drawing_no:
                    image_path = DocumentFiller.find_image_for_output(img_dir, drawing_no)

            # 2.5. 品名으로 시도 (도번 없는 경우)
            if not image_path:
                product_name = DocumentFiller.value_by_aliases(
                    nrow, ["品  名", "品 名", "product_name", "품명"])
                if product_name:
                    sanitized_pname = re.sub(DocumentConstants.INVALID_FILENAME_CHARS, DocumentConstants.INVALID_FILENAME_REPLACEMENT, product_name.strip())
                    image_path = DocumentFiller.find_image_for_output(img_dir, sanitized_pname)

            # 3. out_name(연번)으로 시도
            if not image_path:
                image_path = DocumentFiller.find_image_for_output(img_dir, stem)

            # 4. docx 인접 .data/ 검색 (신규 발행 첨부 이미지)
            if not image_path:
                search_name = image_name or stem
                image_path = DocumentFiller.find_image_for_output(
                    docx_path.parent / ".data", search_name)

            DocumentFiller.insert_images_by_token(doc, image_path)
            doc.save(str(docx_path))

            if callback:
                callback(f"✓ Word 파일 재생성 완료: {docx_path.name}")
            return True
        except Exception as e:
            if callback:
                callback(f"Word 재생성 오류: {e}")
            return False


# ============================================================================
# MoldHistoryCard - 데이터 모델 및 검증 레이어
# ============================================================================
@dataclass
class MoldHistoryCard:
    """금형 이력카드 데이터 모델.

    Dict[str, str] 형태의 XLSX 행 데이터를 구조화된 객체로 감싸고
    기본적인 유효성 검증(validate)을 제공한다.
    신규 발행(NewCardManager) 및 동기화(DocxSyncManager)에서 활용한다.
    """
    # 식별 필드 (필수)
    file_name: str
    management_no: str      # 管理番号
    product_name: str       # 品名
    drawing_no: str         # 図番番号

    # 선택 필드
    storage_company: Optional[str] = None   # 保管会社名
    created_date: Optional[str] = None      # 作成日子
    classification: Optional[str] = None    # 分 類
    current_storage: Optional[str] = None  # 現 保管処
    maker: Optional[str] = None             # 製作処
    model_name: Optional[str] = None        # MODEL名
    mass_prod: Optional[str] = None         # 量産処
    mold_size: Optional[str] = None         # 金型規格
    material_base: Optional[str] = None     # 金型材質-BASE
    material_core: Optional[str] = None     # 金型材質-CORE
    cavity_count: Optional[str] = None      # CAVITY 数
    mold_lifetime: Optional[str] = None     # 金型寿命
    gate_type: Optional[str] = None         # GATE 型式
    machine: Optional[str] = None           # 使用機械
    contract_date: Optional[str] = None     # 契約日
    approval_date: Optional[str] = None     # 承認日
    mold_price: Optional[str] = None        # 金型価
    image_filename: Optional[str] = None    # 金型写真

    # 체크박스 필드
    is_new: bool = False          # 新作
    is_increase: bool = False     # 増作
    is_dual: bool = False         # 二元化
    is_vendor_change: bool = False  # 業者変更
    is_spec_change: bool = False    # 仕様変更

    def validate(self) -> List[str]:
        """필수 필드 및 형식 검증. 오류 메시지 목록을 반환한다 (빈 목록 = 정상)."""
        errors: List[str] = []

        if not self.management_no:
            errors.append("管理番号가 비어있습니다")
        if not self.product_name:
            errors.append("品名이 비어있습니다")
        if not self.drawing_no:
            errors.append("図番番号가 비어있습니다")

        # 연번 형식 검증 (예: 19-001)
        if self.file_name and not re.match(r"^\d{2}-\d{3}$", self.file_name):
            errors.append(f"파일명 형식 오류: {self.file_name!r} (예: 19-001)")

        # 날짜 형식 검증 (YYYY.MM.DD)
        if self.created_date:
            try:
                datetime.strptime(self.created_date, "%Y.%m.%d")
            except ValueError:
                errors.append(f"작성일자 형식 오류: {self.created_date!r} (예: 2024.01.15)")

        return errors

    @classmethod
    def from_dict(cls, data: Dict[str, str]) -> 'MoldHistoryCard':
        """XLSX 행 딕셔너리에서 MoldHistoryCard 생성"""
        def _bool(key: str) -> bool:
            return data.get(key, "").strip() == "1"

        return cls(
            file_name=data.get("File name", ""),
            management_no=data.get("管理番号", ""),
            product_name=data.get("品 名", ""),
            drawing_no=data.get("図番番号", ""),
            storage_company=data.get("保管会社名") or None,
            created_date=data.get("作成日子") or None,
            classification=data.get("分 類") or None,
            current_storage=data.get("現 保管処") or None,
            maker=data.get("製作処") or None,
            model_name=data.get("MODEL名") or None,
            mass_prod=data.get("量産処") or None,
            mold_size=data.get("金型規格") or None,
            material_base=data.get("金型材質-BASE") or None,
            material_core=data.get("金型材質-CORE") or None,
            cavity_count=data.get("CAVITY 数") or None,
            mold_lifetime=data.get("金型寿命") or None,
            gate_type=data.get("GATE 型式") or None,
            machine=data.get("使用機械") or None,
            contract_date=data.get("契約日") or None,
            approval_date=data.get("承認日") or None,
            mold_price=data.get("金型価") or None,
            image_filename=data.get("金型写真") or None,
            is_new=_bool("新作"),
            is_increase=_bool("増作"),
            is_dual=_bool("二元化"),
            is_vendor_change=_bool("業者変更"),
            is_spec_change=_bool("仕様変更"),
        )

    def to_dict(self) -> Dict[str, str]:
        """XLSX 행 딕셔너리로 변환 (DocumentFiller 호환)"""
        def _flag(val: bool) -> str:
            return "1" if val else ""

        return {
            "File name": self.file_name,
            "管理番号": self.management_no,
            "品 名": self.product_name,
            "図番番号": self.drawing_no,
            "保管会社名": self.storage_company or "",
            "作成日子": self.created_date or "",
            "分 類": self.classification or "",
            "現 保管処": self.current_storage or "",
            "製作処": self.maker or "",
            "MODEL名": self.model_name or "",
            "量産処": self.mass_prod or "",
            "金型規格": self.mold_size or "",
            "金型材質-BASE": self.material_base or "",
            "金型材質-CORE": self.material_core or "",
            "CAVITY 数": self.cavity_count or "",
            "金型寿命": self.mold_lifetime or "",
            "GATE 型式": self.gate_type or "",
            "使用機械": self.machine or "",
            "契約日": self.contract_date or "",
            "承認日": self.approval_date or "",
            "金型価": self.mold_price or "",
            "新作": _flag(self.is_new),
            "増作": _flag(self.is_increase),
            "二元化": _flag(self.is_dual),
            "業者変更": _flag(self.is_vendor_change),
            "仕様変更": _flag(self.is_spec_change),
            "金型写真": self.image_filename or "",
        }


# ============================================================================
# NewCardManager - 신규 이력카드 발행 (기능 4)
# ============================================================================
class NewCardManager:
    """신규 이력카드 발행: 자동 연번 생성 + XLSX 추가 + Word 파일 생성"""

    @classmethod
    def get_next_file_name(cls, xlsx_path: Path) -> str:
        """기존 XLSX에서 최대 연번을 찾아 다음 파일명 반환 (예: 19-006)"""
        rows = DocumentFiller.load_rows_from_xlsx(xlsx_path)
        max_serial = 0
        prefix = ""

        for row in rows:
            file_name = row.get("File name", "").strip()
            m = re.match(r"^(\d{2}-)(\d{3})$", file_name)
            if m:
                serial = int(m.group(2))
                if serial > max_serial:
                    max_serial = serial
                    prefix = m.group(1)

        if not prefix:
            prefix = date.today().strftime("%y") + "-"

        return f"{prefix}{max_serial + 1:03d}"

    @classmethod
    def get_last_entry(cls, xlsx_path: Path) -> Tuple[str, str]:
        """DB 마지막 행의 파일명과 품명 반환 (신규 발행 다이얼로그 표시용)"""
        rows = DocumentFiller.load_rows_from_xlsx(xlsx_path)
        if not rows:
            return "", ""
        last = rows[-1]
        file_name = last.get("File name", "").strip()
        nrow = DocumentFiller.row_norm_map(last)
        product = DocumentFiller.value_by_aliases(nrow, ["品  名", "品 名", "품명"])
        return file_name, (product.strip() if product else "")

    @classmethod
    def sanitize(cls, text: str) -> str:
        return re.sub(DocumentConstants.INVALID_FILENAME_CHARS, DocumentConstants.INVALID_FILENAME_REPLACEMENT, text)

    @classmethod
    def add_to_xlsx(cls, xlsx_path: Path, row_dict: Dict[str, str]) -> None:
        """XLSX 맨 아래에 새 행 추가"""
        wb = load_workbook(xlsx_path)
        ws = wb.active
        headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
        new_row = [row_dict.get(h, "") for h in headers]
        ws.append(new_row)
        wb.save(xlsx_path)
        wb.close()

    @classmethod
    def generate_card(cls, xlsx_path: Path, template_path: Path, output_dir: Path,
                      img_dir: Path, row_dict: Dict[str, str],
                      image_source_path: Optional[Path] = None,
                      callback: Optional[Callable[[str], None]] = None) -> Optional[Path]:
        """신규 이력카드 생성: XLSX 추가 → 이미지 복사(.data/) → Word 파일 생성"""
        try:
            product = cls.sanitize(row_dict.get("品 名", "").strip())
            drawing = cls.sanitize(row_dict.get("図番番号", "").strip())
            image_stem = f"{product}_{drawing}" if (product or drawing) else "image"
            row_dict.setdefault("金型写真", image_stem)

            # 첨부 이미지 처리: .data/ 에 품명_도번.ext 로 복사
            inserted_image_path: Optional[Path] = None
            if image_source_path and image_source_path.exists():
                data_dir = output_dir / ".data"
                data_dir.mkdir(parents=True, exist_ok=True)
                target_name = f"{image_stem}{image_source_path.suffix.lower()}"
                target_path = data_dir / target_name
                shutil.copy2(str(image_source_path), str(target_path))
                row_dict["金型写真"] = image_stem  # 확장자 제외한 스템
                inserted_image_path = target_path
                if callback:
                    callback(f"이미지 복사됨: .data/{target_name}")

            cls.add_to_xlsx(xlsx_path, row_dict)
            if callback:
                callback("XLSX에 새 행 추가 완료")

            file_name = row_dict.get("File name", "new_card").strip()

            doc = Document(str(template_path))
            nrow = DocumentFiller.row_norm_map(row_dict)
            DocumentFiller.replace_in_doc(doc, nrow)

            # 이미지 탐색: 직접 경로 → img_dir → .data/
            if inserted_image_path is None:
                search = row_dict.get("金型写真", "")
                if search:
                    inserted_image_path = DocumentFiller.find_image_for_output(img_dir, search)
                if not inserted_image_path and search:
                    inserted_image_path = DocumentFiller.find_image_for_output(
                        output_dir / ".data", search
                    )

            DocumentFiller.insert_images_by_token(doc, inserted_image_path)

            output_dir.mkdir(parents=True, exist_ok=True)
            out_path = DocumentFiller.unique_path(output_dir / f"{file_name}.docx")
            doc.save(str(out_path))

            if callback:
                callback(f"✓ 신규 이력카드 생성 완료: {out_path.name}")
            return out_path
        except Exception as e:
            if callback:
                callback(f"✗ 생성 오류: {e}")
            return None


if __name__ == "__main__":
    # CLI 사용 예제
    parser = argparse.ArgumentParser(description="Integrated HWP/XLSX processing tool")
    subparsers = parser.add_subparsers(dest='command', help='Commands')

    # HWP → XLSX
    hwp_parser = subparsers.add_parser('hwp2xlsx', help='Convert HWP files to XLSX')
    hwp_parser.add_argument('--input-dir', default='YES', help='Input directory with HWP files')
    hwp_parser.add_argument('--output', default='output_from_hwp.xlsx', help='Output XLSX path')

    # Extract images
    img_parser = subparsers.add_parser('extract-images', help='Extract images from HWP files')
    img_parser.add_argument('--input-dir', default='YES', help='Input directory with HWP files')
    img_parser.add_argument('--output-dir', default='img', help='Output directory for images')

    # XLSX → DOCX
    docx_parser = subparsers.add_parser('xlsx2docx', help='Fill DOCX files from XLSX')
    docx_parser.add_argument('--xlsx', default='output_from_hwp.xlsx', help='Input XLSX path')
    docx_parser.add_argument('--template', default='Word_양식.docx', help='Template DOCX path')
    docx_parser.add_argument('--out-dir', default='output', help='Output directory')
    docx_parser.add_argument('--img-dir', default='img', help='Image directory')
    docx_parser.add_argument('--limit', type=int, default=0, help='Process first N rows only')

    args = parser.parse_args()

    if args.command == 'hwp2xlsx':
        HWPProcessor.process(Path(args.input_dir), Path(args.output))
    elif args.command == 'extract-images':
        input_dir = Path(args.input_dir)
        hwp_files = sorted(input_dir.glob('*.hwp'))
        total_extracted = 0
        for hwp_file in hwp_files:
            extractor = HWPImageExtractor(str(hwp_file), args.output_dir)
            extracted = extractor.extract_images()
            if extracted > 0:
                print(f"[+] {hwp_file.stem}: {extracted} image(s)")
                total_extracted += extracted
        print(f"\n[*] Total extracted: {total_extracted}")
    elif args.command == 'xlsx2docx':
        DocumentFiller.process(
            Path(args.xlsx),
            Path(args.template),
            Path(args.out_dir),
            Path(args.img_dir),
            args.limit
        )
    elif args.command == 'docx2pdf':
        from pdf import DocxToPdfConverter
        converter = DocxToPdfConverter()
        converter.convert(Path(args.input), Path(args.output))
    elif args.command == 'merge-docx-pdf':
        from pdf import DocxToPdfConverter
        converter = DocxToPdfConverter()
        input_files = [Path(f) for f in args.inputs]
        output_path = Path(args.output)
        converter.convert_and_merge(input_files, output_path)

